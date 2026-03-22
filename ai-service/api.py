"""
api.py
======
FastAPI server untuk APILA AI Service
Mendukung: Chat, OCR (gambar), PDF parsing, Word parsing

Cara pakai:
    pip install fastapi uvicorn python-multipart pillow pdfplumber python-docx
    python api.py

Akses di: http://localhost:8000
"""

from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from typing import Optional, List
import sys
from pathlib import Path
import io
import os

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent))

from dotenv import load_dotenv
from openai import OpenAI
from retriever import retrieve_docs, format_for_prompt

load_dotenv()

app = FastAPI(
    title="APILA API",
    description="AI-Powered Indonesian Law Assistant API - with OCR and Document Parsing",
    version="2.0.0"
)

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class ChatRequest(BaseModel):
    message: str
    history: Optional[List[dict]] = []
    extracted_text: Optional[str] = None


class ChatResponse(BaseModel):
    content: str
    sources: List[dict]


DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
DEEPSEEK_MODEL = os.getenv("DEEPSEEK_MODEL", "deepseek-chat")
AI_SERVICE_PORT = int(os.getenv("AI_SERVICE_PORT", "8001"))


def ask_deepseek(question: str, context: str, extracted_text: str = "") -> str:
    """
    Generate jawaban dengan DeepSeek (OpenAI-compatible API).
    """
    if not DEEPSEEK_API_KEY:
        raise RuntimeError("DEEPSEEK_API_KEY belum diset.")

    client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")

    system_prompt = (
        "Kamu adalah APILA, asisten informasi hukum Indonesia. "
        "Jawab dalam Bahasa Indonesia yang jelas dan ringkas. "
        "Gunakan konteks hukum yang diberikan sebagai sumber utama. "
        "Jika konteks kurang, sampaikan keterbatasan dan sarankan konsultasi LBH. "
        "Selalu tambahkan disclaimer bahwa jawaban ini bukan nasihat hukum resmi."
    )

    user_prompt = (
        f"Pertanyaan pengguna:\n{question}\n\n"
        f"Konteks pasal/peraturan:\n{context}\n\n"
    )

    if extracted_text:
        user_prompt += f"Teks dokumen yang diunggah pengguna:\n{extracted_text[:5000]}\n\n"

    completion = client.chat.completions.create(
        model=DEEPSEEK_MODEL,
        temperature=0.1,
        timeout=20,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    )

    content = completion.choices[0].message.content
    if not content:
        raise RuntimeError("DeepSeek mengembalikan respons kosong.")
    return content.strip()


@app.get("/")
def root():
    return {
        "name": "APILA API v2.0",
        "version": "2.0.0",
        "description": "AI-Powered Indonesian Law Assistant with OCR & Document Parsing"
    }


@app.get("/health")
def health_check():
    return {"status": "healthy"}


@app.post("/chat", response_model=ChatResponse)
def chat(request: ChatRequest):
    """
    Endpoint utama untuk chat dengan AI
    """
    try:
        # Get the message - either from text or extracted from document
        user_message = request.message
        context_text = request.extracted_text or ""
        
        # Combine message with extracted text if any
        if context_text:
            full_query = f"{user_message}\n\nDokumen yang diupload:\n{context_text}"
        else:
            full_query = user_message
            
        # Retrieve relevant legal documents
        docs = retrieve_docs(full_query, top_k=5)
        
        if not docs:
            return ChatResponse(
                content="Maaf, saya tidak menemukan informasi hukum yang relevan dengan pertanyaan Anda. Silakan coba pertanyaan lain atau konsultasikan dengan pihak yang lebih kompeten.",
                sources=[]
            )
        
        # Format documents as context
        context = format_for_prompt(docs)

        # Generate response with DeepSeek (fallback ke template jika API gagal)
        try:
            response_content = ask_deepseek(user_message, context, context_text)
        except Exception as err:
            print(f"DeepSeek call failed, using fallback: {err}")
            response_content = generate_response(user_message, context, docs, context_text)
        
        # Format sources for frontend
        sources = [
            {
                "title": f"{doc['judul']} - {doc['pasal']}",
                "snippet": doc['isi'][:200] + "..." if len(doc['isi']) > 200 else doc['isi']
            }
            for doc in docs[:3]
        ]
        
        return ChatResponse(
            content=response_content,
            sources=sources
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/process-document")
async def process_document(
    file: UploadFile = File(...),
    question: Optional[str] = Form(None)
):
    """
    Endpoint untuk memproses dokumen (gambar, PDF, Word)
    dan mengekstrak teks untuk analisis lebih lanjut
    """
    extracted_text = ""
    file_type = file.content_type
    
    try:
        # Read file content
        content = await file.read()
        
        if file_type.startswith('image/'):
            # Process image with OCR
            extracted_text = process_image_ocr(content)
        elif file_type == 'application/pdf':
            # Process PDF
            extracted_text = process_pdf(content)
        elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                          'application/msword']:
            # Process Word document
            extracted_text = process_word(content)
        else:
            raise HTTPException(status_code=400, detail=f"Tipe file tidak didukung: {file_type}")
        
        # If there's a question, also get AI response
        if question:
            docs = retrieve_docs(f"{question}\n\nDokumen:\n{extracted_text}", top_k=5)
            context = format_for_prompt(docs) if docs else ""
            try:
                response = ask_deepseek(question, context, extracted_text)
            except Exception as err:
                print(f"DeepSeek call failed, using fallback: {err}")
                response = generate_response(question, context, docs, extracted_text)
            
            sources = [
                {
                    "title": f"{doc['judul']} - {doc['pasal']}",
                    "snippet": doc['isi'][:200] + "..." if len(doc['isi']) > 200 else doc['isi']
                }
                for doc in docs[:3]
            ]
            
            return JSONResponse({
                "status": "success",
                "extracted_text": extracted_text[:1000] + "..." if len(extracted_text) > 1000 else extracted_text,
                "response": response,
                "sources": sources
            })
        
        return JSONResponse({
            "status": "success",
            "extracted_text": extracted_text,
            "word_count": len(extracted_text.split()),
            "char_count": len(extracted_text)
        })
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")


def process_image_ocr(content: bytes) -> str:
    """
    Process image and extract text using OCR
    Note: Requires pytesseract and tesseract installed
    """
    try:
        from PIL import Image
        
        # Open image from bytes
        image = Image.open(io.BytesIO(content))
        
        # Try to use pytesseract if available
        try:
            import pytesseract
            text = pytesseract.image_to_string(image, lang='ind+eng')
            return text.strip()
        except ImportError:
            # Fallback: return basic image info
            return f"[Gambar: {image.size[0]}x{image.size[1]} piksel - {image.mode}] - Untuk ekstraksi teks penuh, harap install Tesseract OCR"
            
    except Exception as e:
        return f"Error memproses gambar: {str(e)}"


def process_pdf(content: bytes) -> str:
    """
    Extract text from PDF
    """
    try:
        import pdfplumber
        
        text_parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        return "\n\n".join(text_parts)
        
    except ImportError:
        return "Error: pdfplumber tidak terinstall. Install dengan: pip install pdfplumber"
    except Exception as e:
        return f"Error memproses PDF: {str(e)}"


def process_word(content: bytes) -> str:
    """
    Extract text from Word document
    """
    try:
        import docx
        
        doc = docx.Document(io.BytesIO(content))
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    paragraphs.append(row_text)
        
        return "\n\n".join(paragraphs)
        
    except ImportError:
        return "Error: python-docx tidak terinstall. Install dengan: pip install python-docx"
    except Exception as e:
        return f"Error memproses Word: {str(e)}"


def generate_response(question: str, context: str, docs: list, extracted_text: str = "") -> str:
    """
    Generate response based on retrieved documents.
    In production, this would call an LLM like DeepSeek.
    """
    if not docs:
        return "Maaf, saya tidak dapat menemukan informasi yang relevan dengan pertanyaan Anda."
    
    # Build response
    response = ""
    
    # If there's extracted text from document, mention it
    if extracted_text:
        response += f"📄 **Analisis Dokumen Anda:**\n\n"
        response += f"Berikut adalah ringkasan dari dokumen yang Anda upload:\n"
        response += f"_{extracted_text[:500]}..."
        if len(extracted_text) > 500:
            response += f"_\n\n---\n\n"
    
    # Get the most relevant document
    main_doc = docs[0]
    
    response += f"**Berdasarkan pencarian hukum:**\n\n"
    response += f"**{main_doc['judul']}**\n"
    response += f"_{main_doc['pasal']}_\n\n"
    response += f"{main_doc['isi'][:800]}"
    
    if len(main_doc['isi']) > 800:
        response += "...\n\n"
    
    if len(docs) > 1:
        response += f"\n\n*Referensi tambahan dari {len(docs)-1} dokumen hukum lainnya.*"
    
    response += f"\n\n**📋 Sumber:** {main_doc['sumber']}"
    
    return response


if __name__ == "__main__":
    import uvicorn
    print("Starting APILA API v2.0...")
    print(f"API documentation: http://localhost:{AI_SERVICE_PORT}/docs")
    print("API docs endpoint ready.")
    print("Document endpoint: POST /process-document")
    if not DEEPSEEK_API_KEY:
        print("WARNING: DEEPSEEK_API_KEY belum diset, mode fallback template aktif.")
    uvicorn.run(app, host="0.0.0.0", port=AI_SERVICE_PORT)
