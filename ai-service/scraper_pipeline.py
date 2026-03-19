"""
scraper_pipeline.py
====================
Konversi file PDF / DOCX / HTML → chunks per pasal → SQLite database

Cara pakai:
    python scraper_pipeline.py

Tidak butuh API key apapun. Python murni.
"""

import os
import re
import json
import sqlite3
import io
import requests
from pathlib import Path
from bs4 import BeautifulSoup

# ── opsional, hanya diimport kalau file ada ──────────────────────────────────
try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    print("⚠️  pdfplumber tidak terinstall — PDF tidak bisa diproses")

try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("⚠️  python-docx tidak terinstall — DOCX tidak bisa diproses")


# =============================================================================
# KONFIGURASI
# =============================================================================

# Folder tempat kamu taruh file PDF / DOCX
RAW_FOLDER = Path("data/raw")

# Path database output
DB_PATH = Path("data/uu_database.db")

# Daftar UU yang mau di-scrape dari URL (opsional)
# Kosongkan list ini kalau kamu hanya pakai file lokal
URL_SOURCES = [
    # Contoh format — ganti dengan URL yang valid
    # {
    #     "id": "uu_8_1999",
    #     "judul": "UU No.8 Tahun 1999",
    #     "tentang": "Perlindungan Konsumen",
    #     "url": "https://peraturan.go.id/...",
    #     "tipe": "html"   # "html", "pdf", atau "docx"
    # },
]


# =============================================================================
# STEP 1 — EXTRACT TEKS DARI BERBAGAI FORMAT
# =============================================================================

def extract_from_pdf(source) -> str:
    """
    Extract teks dari file PDF.
    source bisa berupa: path file (str/Path) atau bytes
    """
    if not PDF_SUPPORT:
        raise ImportError("Install pdfplumber dulu: pip install pdfplumber")

    if isinstance(source, (str, Path)):
        pdf_file = open(source, "rb")
    else:
        pdf_file = io.BytesIO(source)

    text = ""
    with pdfplumber.open(pdf_file) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

    if isinstance(source, (str, Path)):
        pdf_file.close()

    # Deteksi kalau PDF kemungkinan scan (teks terlalu sedikit)
    if len(text.strip()) < 200:
        print("  ⚠️  Teks sangat sedikit — kemungkinan PDF scan.")
        print("     Pertimbangkan OCR dengan Tesseract atau Vision AI.")

    return text


def extract_from_docx(source) -> str:
    """
    Extract teks dari file DOCX.
    source bisa berupa: path file (str/Path) atau bytes
    """
    if not DOCX_SUPPORT:
        raise ImportError("Install python-docx dulu: pip install python-docx")

    if isinstance(source, (str, Path)):
        doc = docx.Document(str(source))
    else:
        doc = docx.Document(io.BytesIO(source))

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    return "\n".join(paragraphs)


def extract_from_html(url: str) -> str:
    """
    Scrape teks dari halaman HTML.
    """
    headers = {"User-Agent": "Mozilla/5.0"}
    response = requests.get(url, headers=headers, timeout=30)
    response.raise_for_status()

    soup = BeautifulSoup(response.text, "html.parser")

    # Hapus elemen yang tidak perlu
    for tag in soup(["nav", "header", "footer", "script",
                     "style", "aside", "advertisement"]):
        tag.decompose()

    # Ambil konten utama
    # Coba cari div konten dulu, fallback ke body
    main = (
        soup.find("div", class_=re.compile(r"content|main|artikel|peraturan", re.I))
        or soup.find("article")
        or soup.find("main")
        or soup.body
    )

    if main:
        return main.get_text(separator="\n", strip=True)
    return soup.get_text(separator="\n", strip=True)


def auto_extract(source, tipe: str = None) -> str:
    """
    Auto-detect format dan extract teks.
    """
    # Deteksi dari ekstensi file kalau tipe tidak disebutkan
    if tipe is None:
        if isinstance(source, (str, Path)):
            ext = Path(source).suffix.lower()
            if ext == ".pdf":
                tipe = "pdf"
            elif ext in (".docx", ".doc"):
                tipe = "docx"
            else:
                tipe = "html"
        else:
            tipe = "pdf"  # default bytes → asumsikan PDF

    if tipe == "pdf":
        return extract_from_pdf(source)
    elif tipe == "docx":
        return extract_from_docx(source)
    elif tipe == "html":
        return extract_from_html(source)
    else:
        raise ValueError(f"Tipe tidak dikenal: {tipe}")


# =============================================================================
# STEP 2 — CHUNKING PER PASAL
# =============================================================================

def chunk_by_pasal(raw_text: str, meta: dict) -> list[dict]:
    """
    Potong teks menjadi chunks per pasal.

    Mengenali pola:
    - "Pasal 1", "Pasal 1A", "PASAL 1"
    - "BAB I", "BAB II"
    - "Ayat (1)" sebagai sub-chunk
    """
    chunks = []
    current_bab = ""
    current_pasal = "Pembukaan / Umum"
    current_isi = []

    # Bersihkan teks dulu
    lines = raw_text.split("\n")
    clean_lines = []
    for line in lines:
        line = line.strip()
        if line:
            clean_lines.append(line)

    pola_bab   = re.compile(r'^(BAB\s+[IVXLC\d]+)', re.IGNORECASE)
    pola_pasal = re.compile(r'^(Pasal\s+\d+[A-Za-z]?)', re.IGNORECASE)

    def save_current():
        """Simpan pasal yang sedang dikumpulkan."""
        isi = " ".join(current_isi).strip()
        if len(isi) > 30:  # skip yang terlalu pendek
            chunks.append({
                "uu_id"   : meta["id"],
                "judul"   : meta["judul"],
                "tentang" : meta["tentang"],
                "tahun"   : meta.get("tahun", ""),
                "bab"     : current_bab,
                "pasal"   : current_pasal,
                "isi"     : isi,
                "sumber"  : meta.get("sumber", "")
            })

    for line in clean_lines:
        if pola_bab.match(line):
            # Simpan pasal sebelumnya
            save_current()
            current_bab = line
            current_pasal = line
            current_isi = []

        elif pola_pasal.match(line):
            # Simpan pasal sebelumnya
            save_current()
            current_pasal = line
            current_isi = []

        else:
            current_isi.append(line)

    # Jangan lupa simpan pasal terakhir
    save_current()

    return chunks


# =============================================================================
# STEP 3 — SIMPAN KE SQLITE
# =============================================================================

def init_db(db_path: Path):
    """Buat tabel kalau belum ada."""
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS peraturan (
            id       INTEGER PRIMARY KEY AUTOINCREMENT,
            uu_id    TEXT NOT NULL,
            judul    TEXT,
            tentang  TEXT,
            tahun    TEXT,
            bab      TEXT,
            pasal    TEXT,
            isi      TEXT NOT NULL,
            sumber   TEXT
        )
    """)
    # Index untuk mempercepat pencarian teks
    cursor.execute("""
        CREATE INDEX IF NOT EXISTS idx_uu_id
        ON peraturan(uu_id)
    """)
    cursor.execute("""
        CREATE INDEX IF NOT EXISTS idx_tentang
        ON peraturan(tentang)
    """)
    conn.commit()
    conn.close()


def save_to_db(chunks: list[dict], db_path: Path):
    """Insert chunks ke database."""
    if not chunks:
        print("  ⚠️  Tidak ada chunks untuk disimpan.")
        return

    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    cursor.executemany("""
        INSERT INTO peraturan (uu_id, judul, tentang, tahun, bab, pasal, isi, sumber)
        VALUES (:uu_id, :judul, :tentang, :tahun, :bab, :pasal, :isi, :sumber)
    """, chunks)

    conn.commit()
    conn.close()
    print(f"  ✅ {len(chunks)} pasal berhasil disimpan")


def delete_uu(uu_id: str, db_path: Path):
    """Hapus UU tertentu dari DB (untuk update)."""
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM peraturan WHERE uu_id = ?", (uu_id,))
    deleted = cursor.rowcount
    conn.commit()
    conn.close()
    print(f"  🗑️  {deleted} baris lama dihapus untuk {uu_id}")


# =============================================================================
# STEP 4 — PIPELINE UTAMA
# =============================================================================

def process_file(filepath: Path, meta: dict = None):
    """
    Proses satu file lokal (PDF atau DOCX).
    """
    if meta is None:
        # Auto-generate meta dari nama file
        stem = filepath.stem  # nama file tanpa ekstensi
        meta = {
            "id"      : stem.lower().replace(" ", "_"),
            "judul"   : stem,
            "tentang" : stem,
            "tahun"   : "",
            "sumber"  : str(filepath)
        }

    print(f"\n📄 Memproses file: {filepath.name}")
    print(f"   ID      : {meta['id']}")
    print(f"   Tentang : {meta['tentang']}")

    # Extract teks
    try:
        raw_text = auto_extract(filepath)
        print(f"   Teks    : {len(raw_text)} karakter diekstrak")
    except Exception as e:
        print(f"   ❌ Gagal extract: {e}")
        return

    # Chunking
    chunks = chunk_by_pasal(raw_text, meta)
    print(f"   Chunks  : {len(chunks)} pasal ditemukan")

    if not chunks:
        print("   ⚠️  Tidak ada pasal yang terdeteksi.")
        print("      Pastikan format teks mengandung 'Pasal X'")
        return

    # Hapus data lama kalau ada, lalu simpan
    delete_uu(meta["id"], DB_PATH)
    save_to_db(chunks, DB_PATH)


def process_url(source_info: dict):
    """
    Proses satu URL (HTML / PDF online).
    """
    print(f"\n🌐 Memproses URL: {source_info['judul']}")
    print(f"   URL: {source_info['url']}")

    try:
        tipe = source_info.get("tipe", "html")

        if tipe in ("pdf", "docx"):
            response = requests.get(
                source_info["url"],
                headers={"User-Agent": "Mozilla/5.0"},
                timeout=30
            )
            response.raise_for_status()
            raw_text = auto_extract(response.content, tipe)
        else:
            raw_text = extract_from_html(source_info["url"])

        print(f"   Teks    : {len(raw_text)} karakter diekstrak")

    except Exception as e:
        print(f"   ❌ Gagal fetch URL: {e}")
        return

    chunks = chunk_by_pasal(raw_text, source_info)
    print(f"   Chunks  : {len(chunks)} pasal ditemukan")

    delete_uu(source_info["id"], DB_PATH)
    save_to_db(chunks, DB_PATH)


def show_summary():
    """Tampilkan ringkasan isi database."""
    if not DB_PATH.exists():
        return

    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    cursor.execute("SELECT COUNT(*) FROM peraturan")
    total = cursor.fetchone()[0]

    cursor.execute("""
        SELECT uu_id, judul, tentang, COUNT(*) as jumlah_pasal
        FROM peraturan
        GROUP BY uu_id
        ORDER BY uu_id
    """)
    rows = cursor.fetchall()
    conn.close()

    print("\n" + "="*55)
    print("📊 RINGKASAN DATABASE")
    print("="*55)
    for row in rows:
        print(f"  [{row[0]}]")
        print(f"   {row[1]} — {row[2]}")
        print(f"   {row[3]} pasal tersimpan")
        print()
    print(f"  TOTAL: {total} pasal dari {len(rows)} UU")
    print("="*55)


# =============================================================================
# JALANKAN PIPELINE
# =============================================================================

if __name__ == "__main__":

    print("="*55)
    print("🚀 SCRAPER PIPELINE — Konversi UU ke Database")
    print("="*55)

    # Buat folder dan database kalau belum ada
    RAW_FOLDER.mkdir(parents=True, exist_ok=True)
    DB_PATH.parent.mkdir(parents=True, exist_ok=True)
    init_db(DB_PATH)
    print(f"✅ Database siap: {DB_PATH}")

    # ── Proses semua file di folder data/raw/ ────────────────────────────────
    files = list(RAW_FOLDER.glob("*.pdf")) + list(RAW_FOLDER.glob("*.docx"))

    if files:
        print(f"\n📁 Ditemukan {len(files)} file di {RAW_FOLDER}/")
        for filepath in sorted(files):
            process_file(filepath)
    else:
        print(f"\n📁 Tidak ada file di {RAW_FOLDER}/")
        print("   Taruh file PDF/DOCX kamu di folder data/raw/ lalu jalankan ulang.")

    # ── Proses URL (kalau ada di URL_SOURCES) ────────────────────────────────
    if URL_SOURCES:
        print(f"\n🌐 Memproses {len(URL_SOURCES)} sumber URL...")
        for source in URL_SOURCES:
            process_url(source)

    # Tampilkan ringkasan
    show_summary()
    print("\n✅ Pipeline selesai! Database siap dipakai retriever.")
